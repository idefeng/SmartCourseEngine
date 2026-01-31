#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
内容生成器 (Content Generator)
==============================

基于知识库自动生成教学课件包：
- 智能课程大纲编排
- 多模态课件生成（文本脚本、音频脚本）
- 高质量题库生成
- Word 文档导出

作者: SmartCourseEngine Team
日期: 2026-01-31
"""

import os
import sys
import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Any

# Rich 美化输出
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.panel import Panel
from rich.prompt import Prompt
from rich.table import Table

console = Console()

# ============================================================================
# API 配置
# ============================================================================
API_KEY = os.getenv("DEEPSEEK_API_KEY", "your-deepseek-api-key-here")
BASE_URL = os.getenv("API_BASE_URL", "https://api.deepseek.com/v1")
MODEL_NAME = os.getenv("MODEL_NAME", "deepseek-chat")

# ============================================================================
# Prompt 模板
# ============================================================================

OUTLINE_PROMPT = """你是一位资深教学设计专家。请根据以下知识点内容，为主题"{topic}"设计一个完整的课程大纲。

【相关知识点】
{knowledge_points}

【要求】
1. 大纲必须包含以下四个部分：导入、核心讲解、案例分析、总结
2. 每个部分需要有2-3个子要点
3. 子要点必须基于提供的知识点内容，不能凭空捏造

请以 JSON 格式返回：
{{
    "title": "课程标题",
    "introduction": {{
        "title": "导入",
        "points": ["要点1", "要点2"]
    }},
    "core_content": {{
        "title": "核心讲解", 
        "points": ["要点1", "要点2", "要点3"]
    }},
    "case_analysis": {{
        "title": "案例分析",
        "points": ["案例要点1", "案例要点2"]
    }},
    "summary": {{
        "title": "总结",
        "points": ["总结要点1", "总结要点2"]
    }}
}}

仅返回 JSON，不要添加任何其他文字。"""

SCRIPT_PROMPT = """你是一位专业的课程讲师。请为以下课程部分撰写详细的讲解词。

【课程主题】{topic}
【当前部分】{section_title}
【要点内容】
{points}

【参考知识点】
{knowledge_context}

【要求】
1. 讲解词约500字，语言生动易懂
2. 在适当位置标注"【画面建议】：..."提示需要展示的教学素材
3. 内容必须基于提供的知识点，不能编造事实

请直接输出讲解词文本。"""

QUIZ_PROMPT = """你是一位专业的命题专家。请根据以下知识点内容，为"{topic}"生成练习题。

【核心知识点】
{knowledge_point}

【要求】
为这个知识点生成3道题：
1. 单选题（4个选项A/B/C/D）
2. 判断题（正确/错误）
3. 案例分析题（基于实际场景）

每道题必须提供"标准答案"和"详细解析"，所有内容必须来源于原始素材。

请以 JSON 格式返回：
{{
    "single_choice": {{
        "question": "题目",
        "options": {{"A": "选项A", "B": "选项B", "C": "选项C", "D": "选项D"}},
        "answer": "A",
        "explanation": "解析..."
    }},
    "true_false": {{
        "question": "判断题题目",
        "answer": true,
        "explanation": "解析..."
    }},
    "case_analysis": {{
        "scenario": "案例场景描述...",
        "question": "问题",
        "answer": "参考答案...",
        "explanation": "解析..."
    }}
}}

仅返回 JSON。"""


# ============================================================================
# 内容生成器主类
# ============================================================================
class ContentGenerator:
    """
    内容生成器
    
    基于知识库自动生成完整的教学课件包。
    """
    
    def __init__(
        self,
        api_key: str = API_KEY,
        base_url: str = BASE_URL,
        model_name: str = MODEL_NAME
    ):
        """初始化内容生成器"""
        self.api_key = api_key
        self.base_url = base_url
        self.model_name = model_name
        
        self._init_llm()
        self._init_knowledge_manager()
    
    def _init_llm(self):
        """初始化大语言模型"""
        try:
            from langchain_openai import ChatOpenAI
            
            self.llm = ChatOpenAI(
                api_key=self.api_key,
                base_url=self.base_url,
                model=self.model_name,
                temperature=0.7
            )
            console.print("[green]✓[/green] LLM 初始化成功")
        except Exception as e:
            console.print(f"[red]✗[/red] LLM 初始化失败: {e}")
            self.llm = None
    
    def _init_knowledge_manager(self):
        """初始化知识管理器"""
        try:
            from knowledge_manager import KnowledgeManager
            
            base_dir = Path(__file__).parent
            self.knowledge_manager = KnowledgeManager(
                db_path=str(base_dir / "chroma_db")
            )
            console.print("[green]✓[/green] 知识库已连接")
        except Exception as e:
            console.print(f"[red]✗[/red] 知识库连接失败: {e}")
            self.knowledge_manager = None
    
    def _call_llm(self, prompt: str) -> str:
        """调用 LLM"""
        if not self.llm:
            raise Exception("LLM 不可用")
        
        response = self.llm.invoke(prompt)
        return response.content.strip()
    
    def _parse_json(self, text: str) -> Optional[Dict]:
        """安全解析 JSON"""
        try:
            # 尝试清理可能的 markdown 代码块
            if "```" in text:
                match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text)
                if match:
                    text = match.group(1)
            
            return json.loads(text)
        except json.JSONDecodeError:
            console.print("[yellow]⚠ JSON 解析失败[/yellow]")
            return None
    
    def retrieve_knowledge(self, topic: str, top_k: int = 10) -> List[Dict]:
        """从知识库检索相关内容"""
        if not self.knowledge_manager:
            return []
        
        results = self.knowledge_manager.query_knowledge(topic, top_k=top_k)
        return results
    
    def generate_outline(self, topic: str, knowledge_points: List[Dict]) -> Optional[Dict]:
        """生成课程大纲"""
        console.print("[blue]→[/blue] 正在生成课程大纲...")
        
        # 格式化知识点
        kp_text = "\n".join([
            f"- {kp.get('content', '')[:200]}"
            for kp in knowledge_points[:5]
        ])
        
        prompt = OUTLINE_PROMPT.format(
            topic=topic,
            knowledge_points=kp_text
        )
        
        response = self._call_llm(prompt)
        outline = self._parse_json(response)
        
        if outline:
            console.print("[green]✓[/green] 课程大纲生成完成")
        
        return outline
    
    def generate_script(
        self,
        topic: str,
        section_title: str,
        points: List[str],
        knowledge_context: str
    ) -> str:
        """生成讲解脚本"""
        prompt = SCRIPT_PROMPT.format(
            topic=topic,
            section_title=section_title,
            points="\n".join([f"- {p}" for p in points]),
            knowledge_context=knowledge_context[:1000]
        )
        
        return self._call_llm(prompt)
    
    def generate_quiz(self, topic: str, knowledge_point: str) -> Optional[Dict]:
        """生成练习题"""
        prompt = QUIZ_PROMPT.format(
            topic=topic,
            knowledge_point=knowledge_point[:800]
        )
        
        response = self._call_llm(prompt)
        return self._parse_json(response)
    
    def generate_courseware(self, topic: str) -> Dict:
        """
        生成完整课件包
        
        Args:
            topic: 课程主题
            
        Returns:
            课件数据字典
        """
        courseware = {
            "topic": topic,
            "generated_at": datetime.now().isoformat(),
            "outline": None,
            "scripts": [],
            "audio_scripts": [],
            "quizzes": []
        }
        
        # 1. 检索相关知识点
        console.print(Panel(f"正在为主题 [bold]{topic}[/bold] 生成课件...", style="blue"))
        
        knowledge_points = self.retrieve_knowledge(topic, top_k=10)
        
        if not knowledge_points:
            console.print("[yellow]⚠ 未找到相关知识点，将使用通用模板[/yellow]")
            knowledge_context = f"主题：{topic}"
        else:
            console.print(f"[green]✓[/green] 检索到 {len(knowledge_points)} 个相关知识点")
            knowledge_context = "\n".join([kp.get('content', '') for kp in knowledge_points[:5]])
        
        # 2. 生成课程大纲
        outline = self.generate_outline(topic, knowledge_points)
        courseware["outline"] = outline
        
        if not outline:
            console.print("[red]✗ 大纲生成失败[/red]")
            return courseware
        
        # 3. 为每个部分生成讲解脚本
        sections = [
            ("introduction", outline.get("introduction", {})),
            ("core_content", outline.get("core_content", {})),
            ("case_analysis", outline.get("case_analysis", {})),
            ("summary", outline.get("summary", {}))
        ]
        
        console.print("[blue]→[/blue] 正在生成讲解脚本...")
        
        for section_key, section_data in sections:
            if not section_data:
                continue
                
            section_title = section_data.get("title", section_key)
            points = section_data.get("points", [])
            
            console.print(f"  [dim]生成: {section_title}[/dim]")
            
            script = self.generate_script(topic, section_title, points, knowledge_context)
            
            courseware["scripts"].append({
                "section": section_title,
                "content": script
            })
            
            # 提取纯文本作为音频脚本（移除画面建议标注）
            audio_script = re.sub(r'【画面建议】[^【]*', '', script)
            courseware["audio_scripts"].append({
                "section": section_title,
                "content": audio_script.strip()
            })
        
        console.print("[green]✓[/green] 讲解脚本生成完成")
        
        # 4. 生成练习题
        console.print("[blue]→[/blue] 正在生成练习题库...")
        
        core_points = outline.get("core_content", {}).get("points", [])[:3]
        
        for i, point in enumerate(core_points):
            console.print(f"  [dim]生成题目 {i+1}/{len(core_points)}[/dim]")
            
            # 找到最相关的知识点作为出题依据
            related_kp = knowledge_points[i] if i < len(knowledge_points) else {}
            kp_content = related_kp.get('content', point)
            
            quiz = self.generate_quiz(topic, kp_content)
            
            if quiz:
                quiz["knowledge_point"] = point
                quiz["source"] = related_kp.get("source", "课程内容")
                courseware["quizzes"].append(quiz)
        
        console.print("[green]✓[/green] 练习题库生成完成")
        
        return courseware
    
    def export_to_word(self, courseware: Dict, output_dir: str = ".") -> str:
        """
        导出为 Word 文档
        
        Args:
            courseware: 课件数据
            output_dir: 输出目录
            
        Returns:
            输出文件路径
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = Document()
        
        topic = courseware.get("topic", "未命名")
        
        # ===== 封面 =====
        title = doc.add_heading(topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        subtitle = doc.add_paragraph("自动生成课件")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        date_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # ===== 目录 =====
        doc.add_heading("目录", level=1)
        doc.add_paragraph("一、课程大纲")
        doc.add_paragraph("二、讲解脚本")
        doc.add_paragraph("三、音频脚本（TTS）")
        doc.add_paragraph("四、练习题库")
        
        doc.add_page_break()
        
        # ===== 一、课程大纲 =====
        doc.add_heading("一、课程大纲", level=1)
        
        outline = courseware.get("outline", {})
        if outline:
            course_title = outline.get("title", topic)
            doc.add_heading(course_title, level=2)
            
            sections = [
                outline.get("introduction", {}),
                outline.get("core_content", {}),
                outline.get("case_analysis", {}),
                outline.get("summary", {})
            ]
            
            for section in sections:
                if not section:
                    continue
                section_title = section.get("title", "")
                doc.add_heading(section_title, level=3)
                
                for point in section.get("points", []):
                    doc.add_paragraph(f"• {point}", style='List Bullet')
        
        doc.add_page_break()
        
        # ===== 二、讲解脚本 =====
        doc.add_heading("二、讲解脚本", level=1)
        
        for script in courseware.get("scripts", []):
            doc.add_heading(script.get("section", ""), level=2)
            doc.add_paragraph(script.get("content", ""))
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # ===== 三、音频脚本 =====
        doc.add_heading("三、音频脚本（TTS用）", level=1)
        
        for script in courseware.get("audio_scripts", []):
            doc.add_heading(script.get("section", ""), level=2)
            doc.add_paragraph(script.get("content", ""))
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # ===== 四、练习题库 =====
        doc.add_heading("四、练习题库", level=1)
        
        for i, quiz in enumerate(courseware.get("quizzes", []), 1):
            doc.add_heading(f"第{i}组题目", level=2)
            doc.add_paragraph(f"【知识点】{quiz.get('knowledge_point', '')}")
            doc.add_paragraph(f"【来源】{quiz.get('source', '')}")
            doc.add_paragraph()
            
            # 单选题
            sc = quiz.get("single_choice", {})
            if sc:
                doc.add_heading("单选题", level=3)
                doc.add_paragraph(sc.get("question", ""))
                
                for key, value in sc.get("options", {}).items():
                    doc.add_paragraph(f"{key}. {value}")
                
                doc.add_paragraph(f"【答案】{sc.get('answer', '')}")
                doc.add_paragraph(f"【解析】{sc.get('explanation', '')}")
                doc.add_paragraph()
            
            # 判断题
            tf = quiz.get("true_false", {})
            if tf:
                doc.add_heading("判断题", level=3)
                doc.add_paragraph(tf.get("question", ""))
                doc.add_paragraph(f"【答案】{'正确' if tf.get('answer') else '错误'}")
                doc.add_paragraph(f"【解析】{tf.get('explanation', '')}")
                doc.add_paragraph()
            
            # 案例分析题
            ca = quiz.get("case_analysis", {})
            if ca:
                doc.add_heading("案例分析题", level=3)
                doc.add_paragraph(f"【案例】{ca.get('scenario', '')}")
                doc.add_paragraph(f"【问题】{ca.get('question', '')}")
                doc.add_paragraph(f"【参考答案】{ca.get('answer', '')}")
                doc.add_paragraph(f"【解析】{ca.get('explanation', '')}")
            
            doc.add_paragraph()
        
        # 保存文件
        safe_topic = re.sub(r'[\\/*?:"<>|]', '_', topic)
        output_path = Path(output_dir) / f"{safe_topic}_自动生成课件.docx"
        doc.save(str(output_path))
        
        return str(output_path)


# ============================================================================
# CLI 入口
# ============================================================================
def main():
    """主入口函数"""
    console.print(Panel("""
╔═══════════════════════════════════════════════════════════╗
║        Content Generator - 智能课件生成器                 ║
║                                                           ║
║  功能：课程大纲 | 讲解脚本 | 音频脚本 | 练习题库           ║
╚═══════════════════════════════════════════════════════════╝
    """, style="blue"))
    
    # 检查 API Key
    if API_KEY == "your-deepseek-api-key-here":
        console.print("[red]✗ 错误：未配置 API_KEY[/red]")
        console.print("[dim]请设置环境变量: $env:DEEPSEEK_API_KEY = 'your-key'[/dim]")
        sys.exit(1)
    
    # 初始化生成器
    generator = ContentGenerator()
    
    # 获取用户输入
    console.print()
    topic = Prompt.ask("[bold blue]请输入课程主题[/bold blue]")
    
    if not topic.strip():
        console.print("[yellow]未输入主题，已取消[/yellow]")
        return
    
    console.print()
    
    try:
        # 生成课件
        courseware = generator.generate_courseware(topic)
        
        # 导出 Word
        console.print("[blue]→[/blue] 正在导出 Word 文档...")
        
        base_dir = Path(__file__).parent
        output_dir = base_dir / "generated_courseware"
        output_dir.mkdir(exist_ok=True)
        
        output_path = generator.export_to_word(courseware, str(output_dir))
        
        console.print()
        console.print(Panel(
            f"[green]✓ 课件生成完成！[/green]\n\n"
            f"[bold]输出文件：[/bold]\n{output_path}",
            title="完成",
            border_style="green"
        ))
        
    except Exception as e:
        console.print(f"[red]✗ 生成失败: {e}[/red]")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
